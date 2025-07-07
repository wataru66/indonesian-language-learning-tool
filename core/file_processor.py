"""File processing module for various file formats"""

import os
from pathlib import Path
from typing import List, Dict, Optional
import chardet
from abc import ABC, abstractmethod


class FileProcessorBase(ABC):
    """Base class for file processors"""
    
    @abstractmethod
    def can_process(self, file_path: Path) -> bool:
        """Check if processor can handle this file"""
        pass
        
    @abstractmethod
    def process(self, file_path: Path) -> str:
        """Process file and return text content"""
        pass


class TextFileProcessor(FileProcessorBase):
    """Process plain text files"""
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.txt', '.text']
        
    def process(self, file_path: Path) -> str:
        """Read text file with encoding detection"""
        # Detect encoding
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            result = chardet.detect(raw_data)
            encoding = result['encoding'] or 'utf-8'
            
        # Read with detected encoding
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            # Fallback to utf-8 with errors ignored
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()


class FileProcessor:
    """Main file processor that delegates to specific processors"""
    
    def __init__(self):
        """Initialize with available processors"""
        self.processors = [
            TextFileProcessor(),
        ]
        
        # Try to import optional processors
        self._init_optional_processors()
        
    def _init_optional_processors(self):
        """Initialize optional processors if libraries are available"""
        # Excel processor
        try:
            from openpyxl import load_workbook
            
            class ExcelFileProcessor(FileProcessorBase):
                def can_process(self, file_path: Path) -> bool:
                    return file_path.suffix.lower() in ['.xlsx', '.xls']
                    
                def process(self, file_path: Path) -> str:
                    workbook = load_workbook(file_path, read_only=True)
                    text_content = []
                    
                    for sheet in workbook.worksheets:
                        for row in sheet.iter_rows(values_only=True):
                            row_text = ' '.join(str(cell) for cell in row if cell)
                            if row_text.strip():
                                text_content.append(row_text)
                                
                    workbook.close()
                    return '\n'.join(text_content)
                    
            self.processors.append(ExcelFileProcessor())
        except ImportError:
            pass
            
        # Word processor
        try:
            from docx import Document
            
            class WordFileProcessor(FileProcessorBase):
                def can_process(self, file_path: Path) -> bool:
                    return file_path.suffix.lower() in ['.docx']
                    
                def process(self, file_path: Path) -> str:
                    doc = Document(file_path)
                    text_content = []
                    
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(paragraph.text)
                            
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = ' '.join(cell.text for cell in row.cells)
                            if row_text.strip():
                                text_content.append(row_text)
                                
                    return '\n'.join(text_content)
                    
            self.processors.append(WordFileProcessor())
        except ImportError:
            pass
            
        # PDF processor
        try:
            import PyPDF2
            
            class PDFFileProcessor(FileProcessorBase):
                def can_process(self, file_path: Path) -> bool:
                    return file_path.suffix.lower() == '.pdf'
                    
                def process(self, file_path: Path) -> str:
                    text_content = []
                    
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        
                        for page_num in range(len(pdf_reader.pages)):
                            page = pdf_reader.pages[page_num]
                            text = page.extract_text()
                            if text.strip():
                                text_content.append(text)
                                
                    return '\n'.join(text_content)
                    
            self.processors.append(PDFFileProcessor())
        except ImportError:
            pass
            
    def process_file(self, file_path: str) -> Dict[str, any]:
        """Process a single file"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        if not path.is_file():
            raise ValueError(f"Not a file: {file_path}")
            
        # Find appropriate processor
        processor = None
        for p in self.processors:
            if p.can_process(path):
                processor = p
                break
                
        if not processor:
            raise ValueError(f"Unsupported file type: {path.suffix}")
            
        # Process file
        try:
            content = processor.process(path)
            return {
                'file_path': str(path),
                'file_name': path.name,
                'file_size': path.stat().st_size,
                'content': content,
                'success': True,
                'error': None
            }
        except Exception as e:
            return {
                'file_path': str(path),
                'file_name': path.name,
                'file_size': path.stat().st_size,
                'content': '',
                'success': False,
                'error': str(e)
            }
            
    def process_files(self, file_paths: List[str]) -> List[Dict[str, any]]:
        """Process multiple files"""
        results = []
        for file_path in file_paths:
            result = self.process_file(file_path)
            results.append(result)
        return results
        
    def process_folder(self, folder_path: str, 
                      recursive: bool = True) -> List[Dict[str, any]]:
        """Process all supported files in a folder"""
        folder = Path(folder_path)
        
        if not folder.exists():
            raise FileNotFoundError(f"Folder not found: {folder_path}")
            
        if not folder.is_dir():
            raise ValueError(f"Not a folder: {folder_path}")
            
        # Get all files
        file_paths = []
        if recursive:
            for ext in self.get_supported_extensions():
                file_paths.extend(folder.rglob(f"*{ext}"))
        else:
            for ext in self.get_supported_extensions():
                file_paths.extend(folder.glob(f"*{ext}"))
                
        # Process files
        return self.process_files([str(p) for p in file_paths])
        
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        extensions = set()
        
        # Check each processor for supported extensions
        for processor in self.processors:
            # Create dummy paths to check
            test_extensions = ['.txt', '.text', '.xlsx', '.xls', 
                             '.docx', '.pdf', '.csv']
            for ext in test_extensions:
                dummy_path = Path(f"dummy{ext}")
                if processor.can_process(dummy_path):
                    extensions.add(ext)
                    
        return sorted(list(extensions))
        
    def combine_contents(self, results: List[Dict[str, any]]) -> str:
        """Combine contents from multiple file results"""
        contents = []
        
        for result in results:
            if result['success'] and result['content']:
                # Add file separator
                contents.append(f"\n--- {result['file_name']} ---\n")
                contents.append(result['content'])
                
        return '\n'.join(contents)